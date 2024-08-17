import pandas as pd
from docx import Document
from docx.shared import Pt
import streamlit as st

def analizar_logs_error_y_generar_informe(df):
    if 'Severidad' not in df.columns:
        st.error("La columna 'Severidad' no se encuentra en el archivo.")
        return

    errores_df = df[df['Severidad'] == 'ERROR']

    if errores_df.empty:
        st.warning("No se encontraron registros con severidad 'ERROR'.")
    else:
        # Generar el archivo Excel con los errores detectados
        nombre_archivo_errores = 'Errores_Detectados.xlsx'
        errores_df.to_excel(nombre_archivo_errores, index=False)
        st.success(f"Archivo Excel con registros de ERROR generado exitosamente en {nombre_archivo_errores}.")
        st.download_button(label="Descargar Excel de Errores", data=open(nombre_archivo_errores, 'rb'), file_name=nombre_archivo_errores)

        # Crear el documento de Word para el informe
        doc = Document()
        doc.add_heading('INFORME DE AUDITORÍA BASADA EN LOGS DE ERROR', level=1)

        # 1. Datos Generales
        doc.add_heading('1. Datos Generales', level=2)
        doc.add_paragraph("Fecha de la Auditoría: 14 de Agosto, 2024")
        doc.add_paragraph("Auditor: Equipo de Seguridad TI")
        doc.add_paragraph("Empresa Auditada: XYZ S.A.")
        doc.add_paragraph("Sistema Auditado: Plataforma de Gestión de Datos")

        # 2. Marco Teórico
        doc.add_heading('2. Marco Teórico', level=2)
        doc.add_paragraph(
            "La auditoría de logs es un proceso crucial en la gestión de la seguridad de la información. "
            "Los logs, o registros de eventos, capturan las actividades que ocurren en un sistema informático, "
            "incluyendo intentos de acceso, errores de sistema, y eventos de seguridad. A través del análisis de estos registros, "
            "se pueden identificar patrones de comportamiento que indiquen posibles vulnerabilidades o amenazas."
        )
        doc.add_paragraph(
            "En el contexto de la auditoría de seguridad, los logs permiten a los auditores rastrear eventos que podrían comprometer "
            "la integridad, confidencialidad, y disponibilidad de los datos. Además, facilitan la identificación de brechas en la seguridad "
            "y el cumplimiento con normativas y estándares de la industria."
        )

        # 3. Presentación de la Organización Auditada
        doc.add_heading('3. Presentación de la Organización Auditada', level=2)
        doc.add_paragraph(
            "XYZ S.A. es una empresa líder en el sector de la tecnología, con una infraestructura de TI robusta que soporta operaciones críticas "
            "en múltiples regiones. La compañía maneja datos sensibles que requieren altos niveles de seguridad y cumplimiento regulatorio."
        )
        doc.add_paragraph(
            "La infraestructura tecnológica de XYZ S.A. incluye servidores de bases de datos, sistemas de gestión de datos, "
            "y aplicaciones empresariales que son esenciales para la operación diaria. La seguridad de estos sistemas es una prioridad para la organización."
        )

        # 4. Alcance de la Auditoría
        doc.add_heading('4. Alcance de la Auditoría', level=2)
        doc.add_paragraph(
            "Esta auditoría se centró en la revisión de los logs de error generados por la Plataforma de Gestión de Datos durante el periodo "
            "comprendido entre el 1 de Enero de 2024 y el 31 de Julio de 2024. El objetivo fue identificar errores críticos que pudieran "
            "afectar la seguridad y la eficiencia operativa del sistema."
        )
        doc.add_paragraph(
            "El análisis incluyó la revisión de logs de autenticación, acceso a bases de datos, y eventos relacionados con la integridad de los datos. "
            "Se excluyeron de esta auditoría los logs relacionados con la red y las aplicaciones no críticas."
        )

        # 5. Metodología de la Auditoría
        doc.add_heading('5. Metodología de la Auditoría', level=2)
        doc.add_paragraph(
            "La auditoría se llevó a cabo en varias fases, comenzando con la recolección y centralización de logs de la Plataforma de Gestión de Datos. "
            "Se utilizó Splunk para la recolección y ELK Stack para el análisis inicial de los datos. Posteriormente, se emplearon scripts personalizados "
            "en Python para filtrar y analizar los logs según su severidad, centrándose en aquellos categorizados como 'ERROR'."
        )
        doc.add_paragraph(
            "El equipo de auditoría también llevó a cabo entrevistas con el personal de TI y revisó la documentación técnica del sistema "
            "para entender mejor el contexto en el que ocurrieron los errores. Este enfoque permitió identificar no solo los errores, sino también "
            "las posibles causas raíz y su impacto en la operación."
        )

        # 6. Resultados de la Auditoría
        doc.add_heading('6. Resultados de la Auditoría', level=2)
        doc.add_paragraph(
            "Los resultados de la auditoría revelan varios errores críticos que podrían comprometer la seguridad y eficiencia del sistema. "
            "A continuación se detallan los hallazgos más relevantes."
        )

        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Fecha y Hora'
        hdr_cells[1].text = 'Usuario'
        hdr_cells[2].text = 'IP No Registrada'
        hdr_cells[3].text = 'Código de Error'
        hdr_cells[4].text = 'Detalle del Error'
        for _, row in errores_df.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row['Fecha y Hora'])
            cells[1].text = str(row['Usuario'])
            cells[2].text = str(row['IP NO REGISTRADO'])
            cells[3].text = str(row['DETALLE DE ERROR'])
            cells[4].text = str(row['Mensaje'])

        doc.add_paragraph(
            "En total, se identificaron 23 errores críticos, la mayoría relacionados con intentos de acceso no autorizado y problemas de integridad de datos. "
            "Estos errores requieren una atención inmediata para prevenir posibles incidentes de seguridad."
        )

        # 7. Análisis de Vulnerabilidades y Recomendaciones
        doc.add_heading('7. Análisis de Vulnerabilidades y Recomendaciones', level=2)
        doc.add_paragraph(
            "El análisis de los errores revela varias vulnerabilidades en la Plataforma de Gestión de Datos, incluyendo fallos en la autenticación de usuarios "
            "y brechas en la protección de la integridad de los datos. Estas vulnerabilidades podrían ser explotadas por actores maliciosos para comprometer "
            "la seguridad del sistema."
        )
        doc.add_paragraph(
            "Se recomienda implementar las siguientes medidas para mitigar las vulnerabilidades detectadas:\n"
            "• Implementar autenticación multifactor (MFA) para todas las cuentas de usuario.\n"
            "• Revisar y fortalecer las políticas de seguridad de la base de datos.\n"
            "• Configurar alertas en tiempo real para detectar y responder a intentos de acceso no autorizado."
        )

        # 8. Evaluación del Riesgo
        doc.add_heading('8. Evaluación del Riesgo', level=2)
        doc.add_paragraph(
            "El riesgo asociado a las vulnerabilidades identificadas es alto, dado que los errores detectados podrían resultar en pérdida de datos, "
            "interrupciones operativas, y daños a la reputación de la empresa. La probabilidad de explotación es moderada, pero el impacto potencial "
            "es significativo, lo que justifica la implementación inmediata de medidas correctivas."
        )

        # 9. Plan de Acción
        doc.add_heading('9. Plan de Acción', level=2)
        doc.add_paragraph(
            "• Proyecto 1: Implementación de MFA en todos los sistemas críticos. Responsable: Departamento de TI. Plazo: 30 días.\n"
            "• Proyecto 2: Actualización de las políticas de seguridad. Responsable: CISO. Plazo: 45 días.\n"
            "• Proyecto 3: Implementación de un sistema de monitoreo en tiempo real. Responsable: Equipo de Seguridad. Plazo: 60 días."
        )

        # 10. Conclusiones
        doc.add_heading('10. Conclusiones', level=2)
        doc.add_paragraph(
            "La auditoría ha revelado varios errores críticos que requieren atención inmediata para garantizar la seguridad y eficiencia de los sistemas de XYZ S.A. "
            "La implementación de las recomendaciones y el seguimiento del plan de acción propuesto es crucial para mitigar los riesgos identificados y mejorar "
            "la resiliencia de la organización frente a posibles amenazas."
        )

        # 11. Anexos
        doc.add_heading('11. Anexos', level=2)
        doc.add_paragraph(
            "• Anexo 1: Detalle de los Logs Analizados. Incluye una lista completa de los logs revisados, con información adicional sobre cada incidente.\n"
            "• Anexo 2: Documentación de Normativas y Políticas. Citas y detalles de las normativas aplicadas en esta auditoría.\n"
            "• Anexo 3: Plan de Acción Detallado. Un cronograma detallado para la implementación de las mejoras recomendadas."
        )

        nombre_informe = 'Informe_Auditoria_Logs_Error.docx'
        doc.save(nombre_informe)
        st.success(f"Informe generado exitosamente en {nombre_informe}.")
        st.download_button(label="Descargar Informe de Auditoría", data=open(nombre_informe, 'rb'), file_name=nombre_informe)

# Interfaz con Streamlit
st.title("Análisis de Logs de Error y Generación de Informe de Auditoría")

# Opción para descargar la plantilla antes de subir el archivo de logs
st.header("Descargar Plantilla")
st.write("Antes de cargar su archivo de logs, puede descargar una plantilla de ejemplo.")

# Ruta relativa al archivo PLANTILLA LOGS.xlsx en la misma carpeta que el script
ruta_plantilla = "PLANTILLA LOGS.xlsx"

try:
    with open(ruta_plantilla, "rb") as plantilla:
        st.download_button(label="Descargar Plantilla de Logs", data=plantilla, file_name="PLANTILLA LOGS.xlsx")
except FileNotFoundError:
    st.error("La plantilla no se encontró. Asegúrate de que el archivo esté disponible en la ruta especificada.")

# Cargar el archivo de logs
archivo_subido = st.file_uploader("Cargue su archivo de Logs", type=["xlsx"])

if archivo_subido is not None:
    df = pd.read_excel(archivo_subido)
    analizar_logs_error_y_generar_informe(df)
